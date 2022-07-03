from docx import Document

def modification(paragraphs,find,modif):
    index=-1
    for i in range(0, len(paragraphs)):
        print(i)
        if paragraphs[i].text.find(find) != -1:
            index=i
            break
    if index!= -1:
       temp= paragraphs[index].text
       pos= temp.find(find)
       p =paragraphs[index].clear()
       temp=temp[:pos]+modif+temp[pos+1:-1]
       run = p.add_run(temp)
       font = run.font
       font.size = 540000


document = Document(r"C:\Users\29545\Desktop\测试.docx")  
print(document.paragraphs[2].text)
print(document.paragraphs[3].text)
print(document.paragraphs[4].text)

modification(document.paragraphs,"color","739639550")